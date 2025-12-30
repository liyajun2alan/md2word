"""
模板格式提取模块

从 Word 模板文件中提取各种元素的格式参数
"""

from docx import Document
from typing import Dict, Any, Optional


class TemplateParser:
    """模板格式提取器"""

    # 标识符映射
    MARKERS = {
        '[一级标题]': 'heading1',
        '[二级标题]': 'heading2',
        '[三级标题]': 'heading3',
        '[四级标题]': 'heading4',
        '[五级标题]': 'heading5',
        '[正文]': 'body',
        '[表格]': 'table',
        '[列表]': 'list',
        '[代码]': 'code',
        '[引用]': 'quote'
    }

    def __init__(self, template_path: str):
        """初始化，加载模板文档

        Args:
            template_path: 模板文件路径
        """
        self.template_doc = Document(template_path)
        self.format_styles = {}

    def extract_all_formats(self) -> Dict[str, Dict[str, Any]]:
        """提取所有标识符的格式

        Returns:
            格式字典，键为元素类型，值为格式参数字典
        """
        # 提取段落格式（通过标识符）
        for para in self.template_doc.paragraphs:
            text = para.text.strip()
            if text in self.MARKERS:
                marker_type = self.MARKERS[text]
                self.format_styles[marker_type] = self._extract_paragraph_format(para)

        # 提取表格格式
        for table in self.template_doc.tables:
            if self._is_marker_table(table):
                self.format_styles['table'] = self._extract_table_format(table)
                break

        # 如果没有提取到表格格式且模版中有表格，使用第一个表格的格式
        if 'table' not in self.format_styles and len(self.template_doc.tables) > 0:
            self.format_styles['table'] = self._extract_table_format(self.template_doc.tables[0])

        # 如果没有提取到标题格式，尝试从 Heading 样式中提取
        if not any(k.startswith('heading') for k in self.format_styles.keys()):
            self._extract_from_heading_styles()

        return self.format_styles

    def _extract_from_heading_styles(self):
        """从模板的 Heading 样式中提取格式（当没有标识符时）"""
        try:
            styles = self.template_doc.styles

            # 提取 Heading 1-5 的格式
            for level in range(1, 6):
                style_name = f'Heading {level}'
                format_key = f'heading{level}'

                try:
                    style = styles[style_name]
                    # 从样式中提取格式
                    format_dict = {}

                    # 字体相关
                    if style.font.name:
                        format_dict['font_name'] = style.font.name
                    if style.font.size:
                        format_dict['font_size'] = style.font.size
                    if style.font.bold is not None:
                        format_dict['bold'] = style.font.bold
                    if style.font.italic is not None:
                        format_dict['italic'] = style.font.italic
                    if style.font.color and style.font.color.rgb:
                        format_dict['color'] = style.font.color.rgb

                    # 段落格式
                    pf = style.paragraph_format
                    if pf.alignment is not None:
                        format_dict['alignment'] = pf.alignment
                    if pf.line_spacing is not None:
                        format_dict['line_spacing'] = pf.line_spacing
                    if pf.space_before is not None:
                        format_dict['space_before'] = pf.space_before
                    if pf.space_after is not None:
                        format_dict['space_after'] = pf.space_after
                    if pf.first_line_indent is not None:
                        format_dict['first_line_indent'] = pf.first_line_indent
                    if pf.left_indent is not None:
                        format_dict['left_indent'] = pf.left_indent
                    if pf.right_indent is not None:
                        format_dict['right_indent'] = pf.right_indent

                    self.format_styles[format_key] = format_dict

                except KeyError:
                    # 样式不存在，跳过
                    pass

            # 尝试提取 Normal 或 Body Text 作为正文格式
            for style_name in ['Body Text', 'Normal']:
                try:
                    style = styles[style_name]
                    format_dict = {}

                    if style.font.name:
                        format_dict['font_name'] = style.font.name
                    if style.font.size:
                        format_dict['font_size'] = style.font.size

                    pf = style.paragraph_format
                    if pf.line_spacing is not None:
                        format_dict['line_spacing'] = pf.line_spacing
                    if pf.space_before is not None:
                        format_dict['space_before'] = pf.space_before
                    if pf.space_after is not None:
                        format_dict['space_after'] = pf.space_after
                    if pf.first_line_indent is not None:
                        format_dict['first_line_indent'] = pf.first_line_indent

                    self.format_styles['body'] = format_dict
                    break  # 找到一个就退出
                except KeyError:
                    pass

        except Exception as e:
            # 提取失败，使用空格式字典
            pass

    def _extract_paragraph_format(self, paragraph) -> Dict[str, Any]:
        """提取段落格式参数

        Args:
            paragraph: python-docx 段落对象

        Returns:
            格式参数字典
        """
        # 获取第一个 run（如果存在）
        run = paragraph.runs[0] if paragraph.runs else None

        format_dict = {
            # 字体相关
            'font_name': run.font.name if run and run.font.name else None,
            'font_size': run.font.size if run and run.font.size else None,
            'bold': run.font.bold if run else False,
            'italic': run.font.italic if run else False,
            'color': run.font.color.rgb if run and run.font.color and run.font.color.rgb else None,

            # 段落格式
            'alignment': paragraph.alignment,
            'space_before': paragraph.paragraph_format.space_before,
            'space_after': paragraph.paragraph_format.space_after,
            'line_spacing': paragraph.paragraph_format.line_spacing,
            'first_line_indent': paragraph.paragraph_format.first_line_indent,
            'left_indent': paragraph.paragraph_format.left_indent,
            'right_indent': paragraph.paragraph_format.right_indent,
        }

        return format_dict

    def _extract_table_format(self, table) -> Dict[str, Any]:
        """提取表格格式参数

        Args:
            table: python-docx 表格对象

        Returns:
            表格格式参数字典
        """
        # 提取表格样式
        first_cell = table.rows[0].cells[0]
        cell_para = first_cell.paragraphs[0] if first_cell.paragraphs else None
        cell_run = cell_para.runs[0] if cell_para and cell_para.runs else None

        # 提取表头格式
        header_indent = cell_para.paragraph_format.first_line_indent if cell_para else None
        header_bold = cell_run.font.bold if cell_run else None
        header_alignment = cell_para.alignment if cell_para else None

        # 提取数据行格式（如果有第二行）
        data_indent = None
        data_bold = None
        data_alignment = None
        if len(table.rows) > 1 and table.rows[1].cells:
            data_cell = table.rows[1].cells[0]
            data_para = data_cell.paragraphs[0] if data_cell.paragraphs else None
            if data_para:
                data_indent = data_para.paragraph_format.first_line_indent
                data_alignment = data_para.alignment
                if data_para.runs:
                    data_bold = data_para.runs[0].font.bold

        # 提取单元格内边距
        from docx.oxml.ns import qn
        cell_margins = {}
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))
        if tblPr is not None:
            tblCellMar = tblPr.find(qn('w:tblCellMar'))
            if tblCellMar is not None:
                for side in ['top', 'left', 'bottom', 'right']:
                    elem = tblCellMar.find(qn(f'w:{side}'))
                    if elem is not None:
                        w_val = elem.get(qn('w:w'))
                        if w_val:
                            cell_margins[side] = int(w_val)

        # 提取列宽（转换为比例）
        column_widths = []
        total_width = 0
        if table.rows and table.columns:
            for col in table.columns:
                column_widths.append(col.width)
                total_width += col.width

        # 计算列宽比例
        column_width_ratios = []
        if total_width > 0:
            column_width_ratios = [w / total_width for w in column_widths]

        return {
            'style': table.style.name if table.style else None,
            'cell_font_name': cell_run.font.name if cell_run and cell_run.font.name else None,
            'cell_font_size': cell_run.font.size if cell_run and cell_run.font.size else None,
            'column_widths': column_widths,
            'column_width_ratios': column_width_ratios,
            'header_indent': header_indent,
            'header_bold': header_bold,
            'header_alignment': header_alignment,
            'data_indent': data_indent,
            'data_bold': data_bold,
            'data_alignment': data_alignment,
            'cell_margins': cell_margins,
        }

    def _is_marker_table(self, table) -> bool:
        """判断表格是否包含 [表格] 标识符

        Args:
            table: python-docx 表格对象

        Returns:
            是否为标记表格
        """
        # 检查表格的第一个单元格是否包含 [表格] 标识符
        if table.rows and table.rows[0].cells:
            first_cell_text = table.rows[0].cells[0].text.strip()
            return first_cell_text == '[表格]'
        return False
