"""
Word 文档生成模块

将结构化元素转换为 Word 文档内容，应用格式
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from typing import List, Dict, Any
import requests
from io import BytesIO
from pathlib import Path
import re


class WordGenerator:
    """Word 文档生成器"""

    def __init__(self, output_path: str, format_styles: Dict[str, Dict[str, Any]], template_path: str = None):
        """初始化

        Args:
            output_path: 输出 Word 文件路径
            format_styles: 格式样式字典
            template_path: 模板文件路径（用于创建新文档时继承样式和编号）
        """
        self.output_path = output_path
        self.format_styles = format_styles
        self.template_path = template_path
        self.doc = None

    def load_target_document(self):
        """加载目标 Word 文档，如不存在则基于模板创建"""
        try:
            # 尝试加载已存在的文档
            self.doc = Document(self.output_path)
        except:
            # 文档不存在，基于模板创建新文档
            if self.template_path and Path(self.template_path).exists():
                try:
                    # 基于模板创建（继承模板的样式和编号配置）
                    self.doc = Document(self.template_path)

                    # 清空模板中的示例内容，只保留样式和编号定义
                    # 删除所有段落
                    for para in list(self.doc.paragraphs):
                        p = para._element
                        p.getparent().remove(p)

                    # 删除所有表格
                    for table in list(self.doc.tables):
                        tbl = table._element
                        tbl.getparent().remove(tbl)

                except:
                    # 模板加载失败，创建空白文档
                    self.doc = Document()
            else:
                # 没有模板路径，创建空白文档
                self.doc = Document()

        # 确保文档有列表编号定义
        self._ensure_list_numbering()

    def append_elements(self, elements: List[Dict[str, Any]]):
        """追加元素到文档

        Args:
            elements: 结构化元素列表
        """
        # 添加分页符（新页开始）
        self.doc.add_page_break()

        for element in elements:
            element_type = element['type']

            if element_type.startswith('heading'):
                self._add_heading(element)
            elif element_type == 'body':
                self._add_paragraph(element)
            elif element_type == 'code':
                self._add_code_block(element)
            elif element_type == 'quote':
                self._add_quote(element)
            elif element_type == 'list':
                self._add_list(element)
            elif element_type == 'table':
                self._add_table(element)
            elif element_type == 'image':
                self._add_image(element)

    def _find_multilevel_num_id(self):
        """查找正确的多级编号定义ID（包含5个级别的）"""
        if not hasattr(self.doc.part, 'numbering_part') or self.doc.part.numbering_part is None:
            return None

        numbering_part = self.doc.part.numbering_part
        numbering_element = numbering_part._element

        # 获取所有 abstractNum
        abstractNums = numbering_element.findall(qn('w:abstractNum'))

        # 查找包含 5 个级别的 abstractNum
        target_abstractNumId = None
        for abstractNum in abstractNums:
            lvls = abstractNum.findall(qn('w:lvl'))
            if len(lvls) >= 5:  # 找到至少有5个级别的
                target_abstractNumId = abstractNum.get(qn('w:abstractNumId'))
                break

        if target_abstractNumId is None:
            # 没找到多级编号，使用最后一个
            nums = numbering_element.findall(qn('w:num'))
            if len(nums) > 0:
                return nums[-1].get(qn('w:numId'))
            return None

        # 找到对应的 num
        nums = numbering_element.findall(qn('w:num'))
        for num in nums:
            abstractNumId_elem = num.find(qn('w:abstractNumId'))
            if abstractNumId_elem is not None:
                if abstractNumId_elem.get(qn('w:val')) == target_abstractNumId:
                    return num.get(qn('w:numId'))

        # 如果没找到对应的 num，使用最后一个
        if len(nums) > 0:
            return nums[-1].get(qn('w:numId'))

        return None

    def _ensure_list_numbering(self):
        """确保文档有列表编号定义（bullet和numbered list）"""
        # 确保有numbering part
        if not hasattr(self.doc.part, 'numbering_part') or self.doc.part.numbering_part is None:
            # 如果没有numbering part，添加一个临时段落使其被创建
            temp_para = self.doc.add_paragraph()
            self._add_numbering_to_paragraph(temp_para, 1)
            # 删除临时段落
            p = temp_para._element
            p.getparent().remove(p)

        numbering_part = self.doc.part.numbering_part
        numbering_element = numbering_part._element

        # 检查是否已有bullet和numbered list的编号定义
        # 我们需要创建两个新的abstractNum：一个用于bullet，一个用于numbered
        abstractNums = numbering_element.findall(qn('w:abstractNum'))
        existing_abstract_ids = [int(an.get(qn('w:abstractNumId'))) for an in abstractNums]

        # 获取下一个可用的abstractNumId
        next_abstract_id = max(existing_abstract_ids) + 1 if existing_abstract_ids else 0

        # 创建bullet list的编号定义
        bullet_abstract_id = next_abstract_id
        bullet_abstract = self._create_bullet_abstract_num(bullet_abstract_id)
        numbering_element.append(bullet_abstract)

        # 创建numbered list的编号定义
        numbered_abstract_id = next_abstract_id + 1
        numbered_abstract = self._create_numbered_abstract_num(numbered_abstract_id)
        numbering_element.append(numbered_abstract)

        # 创建对应的num实例
        nums = numbering_element.findall(qn('w:num'))
        existing_num_ids = [int(n.get(qn('w:numId'))) for n in nums]
        next_num_id = max(existing_num_ids) + 1 if existing_num_ids else 1

        # 为bullet list创建num
        bullet_num_id = next_num_id
        bullet_num = self._create_num_element(bullet_num_id, bullet_abstract_id)
        numbering_element.append(bullet_num)

        # 为numbered list创建num
        numbered_num_id = next_num_id + 1
        numbered_num = self._create_num_element(numbered_num_id, numbered_abstract_id)
        numbering_element.append(numbered_num)

        # 保存numId供后续使用
        self.bullet_num_id = str(bullet_num_id)
        self.numbered_num_id = str(numbered_num_id)

    def _create_bullet_abstract_num(self, abstract_id):
        """创建bullet list的抽象编号定义"""
        abstractNum = OxmlElement('w:abstractNum')
        abstractNum.set(qn('w:abstractNumId'), str(abstract_id))

        # 创建一个级别（级别0）
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')

        # 编号格式：bullet
        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), 'bullet')
        lvl.append(numFmt)

        # 显示文本：使用bullet符号（使用·中点符号）
        lvlText = OxmlElement('w:lvlText')
        lvlText.set(qn('w:val'), '·')
        lvl.append(lvlText)

        # 对齐方式
        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)

        # 段落属性：缩进
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '720')  # 左缩进
        ind.set(qn('w:hanging'), '360')  # 悬挂缩进
        pPr.append(ind)
        lvl.append(pPr)

        # 字体
        rPr = OxmlElement('w:rPr')
        rFonts = OxmlElement('w:rFonts')
        rFonts.set(qn('w:ascii'), 'Symbol')
        rFonts.set(qn('w:hAnsi'), 'Symbol')
        rFonts.set(qn('w:hint'), 'default')
        rPr.append(rFonts)
        lvl.append(rPr)

        abstractNum.append(lvl)
        return abstractNum

    def _create_numbered_abstract_num(self, abstract_id):
        """创建numbered list的抽象编号定义"""
        abstractNum = OxmlElement('w:abstractNum')
        abstractNum.set(qn('w:abstractNumId'), str(abstract_id))

        # 创建一个级别（级别0）
        lvl = OxmlElement('w:lvl')
        lvl.set(qn('w:ilvl'), '0')

        # 编号格式：decimal
        numFmt = OxmlElement('w:numFmt')
        numFmt.set(qn('w:val'), 'decimal')
        lvl.append(numFmt)

        # 显示文本：数字加点
        lvlText = OxmlElement('w:lvlText')
        lvlText.set(qn('w:val'), '%1.')
        lvl.append(lvlText)

        # 起始值
        start = OxmlElement('w:start')
        start.set(qn('w:val'), '1')
        lvl.append(start)

        # 对齐方式
        lvlJc = OxmlElement('w:lvlJc')
        lvlJc.set(qn('w:val'), 'left')
        lvl.append(lvlJc)

        # 段落属性：缩进
        pPr = OxmlElement('w:pPr')
        ind = OxmlElement('w:ind')
        ind.set(qn('w:left'), '720')  # 左缩进
        ind.set(qn('w:hanging'), '360')  # 悬挂缩进
        pPr.append(ind)
        lvl.append(pPr)

        abstractNum.append(lvl)
        return abstractNum

    def _create_num_element(self, num_id, abstract_num_id):
        """创建num元素"""
        num = OxmlElement('w:num')
        num.set(qn('w:numId'), str(num_id))

        abstractNumId = OxmlElement('w:abstractNumId')
        abstractNumId.set(qn('w:val'), str(abstract_num_id))
        num.append(abstractNumId)

        return num

    def _add_numbering_to_paragraph(self, paragraph, level):
        """为段落添加编号配置

        Args:
            paragraph: 段落对象
            level: 标题级别 (1-5)
        """
        # 获取正确的编号ID
        num_id = self._find_multilevel_num_id()
        if num_id is None:
            return

        # 为段落添加编号配置
        pPr = paragraph._element.pPr
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            paragraph._element.insert(0, pPr)

        # 删除旧的 numPr（如果存在）
        old_numPr = pPr.find(qn('w:numPr'))
        if old_numPr is not None:
            pPr.remove(old_numPr)

        # 创建新的 numPr
        numPr = OxmlElement('w:numPr')

        # 设置 ilvl（编号级别，0-based）
        ilvl_elem = OxmlElement('w:ilvl')
        ilvl_elem.set(qn('w:val'), str(level - 1))  # level 1 -> ilvl 0
        numPr.append(ilvl_elem)

        # 设置 numId（编号ID）
        numId_elem = OxmlElement('w:numId')
        numId_elem.set(qn('w:val'), num_id)
        numPr.append(numId_elem)

        # 添加到段落属性
        pPr.append(numPr)

    def _add_list_numbering(self, paragraph, is_ordered):
        """为列表段落添加编号配置

        Args:
            paragraph: 段落对象
            is_ordered: 是否为有序列表
        """
        # 使用预先创建的bullet或numbered list的numId
        num_id = self.numbered_num_id if is_ordered else self.bullet_num_id

        # 为段落添加编号配置
        pPr = paragraph._element.pPr
        if pPr is None:
            pPr = OxmlElement('w:pPr')
            paragraph._element.insert(0, pPr)

        # 删除旧的 numPr（如果存在）
        old_numPr = pPr.find(qn('w:numPr'))
        if old_numPr is not None:
            pPr.remove(old_numPr)

        # 创建新的 numPr
        numPr = OxmlElement('w:numPr')

        # 设置 ilvl（编号级别，0-based）
        ilvl_elem = OxmlElement('w:ilvl')
        ilvl_elem.set(qn('w:val'), '0')  # 列表项使用级别0
        numPr.append(ilvl_elem)

        # 设置 numId（编号ID）
        numId_elem = OxmlElement('w:numId')
        numId_elem.set(qn('w:val'), num_id)
        numPr.append(numId_elem)

        # 添加到段落属性
        pPr.append(numPr)

    def _remove_numbering_prefix(self, text: str) -> str:
        """移除标题文本中的编号前缀

        Args:
            text: 原始标题文本

        Returns:
            清理后的标题文本
        """
        # 移除模式：
        # 1. "第X章" "第X节" "第X部分" 等
        # 2. "1."、"1.1"、"1.1.1" 等数字编号
        # 3. "1)"、"(1)" 等括号编号
        # 4. "一、"、"（一）" 等中文编号

        patterns = [
            r'^第[零一二三四五六七八九十百千万\d]+[章节部分条款项][\s　]*',  # 第X章
            r'^[\d]+\.[\d\.]*[\s　]+',  # 1. 或 1.1 或 1.1.1
            r'^[\d]+[)）][\s　]+',  # 1) 或 1）
            r'^\([\d]+\)[\s　]+',  # (1)
            r'^[一二三四五六七八九十]+[、．.][\s　]+',  # 一、
            r'^[\(（][一二三四五六七八九十]+[\)）][\s　]+',  # （一）
            r'^[⑴⑵⑶⑷⑸⑹⑺⑻⑼⑽⑾⑿⒀⒁⒂⒃⒄⒅⒆⒇][\s　]+',  # 带圈数字
        ]

        result = text
        for pattern in patterns:
            result = re.sub(pattern, '', result)

        return result.strip()

    def _add_heading(self, element: Dict[str, Any]):
        """添加标题

        Args:
            element: 标题元素
        """
        level = element['level']
        style_key = element['type']  # heading1, heading2, ...

        # 清理标题文本中的编号前缀（因为Word会自动添加编号）
        cleaned_text = self._remove_numbering_prefix(element['text'])

        # 使用 Word 内置标题样式（支持自动编号）
        # 映射：heading1 -> Heading 1, heading2 -> Heading 2, ...
        word_style = f'Heading {level}'

        try:
            para = self.doc.add_paragraph(cleaned_text, style=word_style)
        except:
            # 如果样式不存在，使用普通段落
            para = self.doc.add_paragraph(cleaned_text)

        # 添加编号配置
        self._add_numbering_to_paragraph(para, level)

        # 应用模板中定义的格式（覆盖样式的默认格式）
        if style_key in self.format_styles:
            self._apply_paragraph_format(para, self.format_styles[style_key])

    def _add_paragraph(self, element: Dict[str, Any]):
        """添加正文段落

        Args:
            element: 段落元素
        """
        text = element.get('text', '').strip()

        # 过滤掉只包含项目符号的空段落
        if text in ['•', '●', '■', '□', '◆', '◇', '▪', '▫']:
            return

        para = self.doc.add_paragraph(element['text'])

        if 'body' in self.format_styles:
            self._apply_paragraph_format(para, self.format_styles['body'])

    def _add_code_block(self, element: Dict[str, Any]):
        """添加代码块

        Args:
            element: 代码块元素
        """
        para = self.doc.add_paragraph(element['text'])

        if 'code' in self.format_styles:
            self._apply_paragraph_format(para, self.format_styles['code'])
        else:
            # 默认代码样式：等宽字体
            for run in para.runs:
                run.font.name = 'Consolas'
                run._element.rPr.rFonts.set(qn('w:eastAsia'), 'Consolas')
                run.font.size = Pt(9)

    def _add_quote(self, element: Dict[str, Any]):
        """添加引用块

        Args:
            element: 引用块元素
        """
        para = self.doc.add_paragraph(element['text'])

        if 'quote' in self.format_styles:
            self._apply_paragraph_format(para, self.format_styles['quote'])
        else:
            # 默认引用样式：缩进 + 灰色
            para.paragraph_format.left_indent = Inches(0.5)
            for run in para.runs:
                run.font.italic = True
                run.font.color.rgb = RGBColor(128, 128, 128)

    def _add_list(self, element: Dict[str, Any]):
        """添加列表

        Args:
            element: 列表元素
        """
        for item in element['items']:
            # 过滤掉只包含项目符号或空白的列表项
            item_text = item.strip()
            if not item_text or item_text in ['•', '●', '■', '□', '◆', '◇', '▪', '▫']:
                continue

            # 创建普通段落
            para = self.doc.add_paragraph(item)

            # 添加列表编号
            self._add_list_numbering(para, element['ordered'])

            # 如果有自定义列表格式，应用它
            if 'list' in self.format_styles:
                self._apply_paragraph_format(para, self.format_styles['list'])


    def _add_table(self, element: Dict[str, Any]):
        """添加表格

        Args:
            element: 表格元素
        """
        num_cols = len(element['header']) if element['header'] else 0
        num_rows = len(element['body']) + 1  # +1 for header

        if num_cols == 0:
            return

        table = self.doc.add_table(rows=num_rows, cols=num_cols)

        # 设置autofit（按内容自动调整列宽）
        table.autofit = True

        # 应用表格样式
        table_format = self.format_styles.get('table', {})

        # 设置表格样式，优先使用模板样式，否则使用 Table Grid（有边框）
        try:
            if table_format.get('style'):
                table.style = table_format['style']
            else:
                table.style = 'Table Grid'
        except:
            # 样式不存在，使用默认
            table.style = 'Table Grid'

        # 应用单元格内边距
        self._apply_table_cell_margins(table, table_format)

        # 如果模版中没有表格格式，设置默认格式
        if not table_format:
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            table_format = {
                'cell_font_name': '宋体',
                'cell_font_size': Pt(10.5),
                'header_indent': None,
                'header_bold': True,
                'header_alignment': WD_ALIGN_PARAGRAPH.CENTER,
                'data_indent': None,
                'data_bold': None,
                'data_alignment': WD_ALIGN_PARAGRAPH.LEFT,
            }
        else:
            # 补充缺失的默认值
            from docx.enum.text import WD_ALIGN_PARAGRAPH
            table_format.setdefault('cell_font_name', '宋体')
            table_format.setdefault('cell_font_size', Pt(10.5))
            table_format.setdefault('header_bold', True)
            table_format.setdefault('header_alignment', WD_ALIGN_PARAGRAPH.CENTER)
            table_format.setdefault('data_alignment', WD_ALIGN_PARAGRAPH.LEFT)

        # 填充表头
        for i, cell_text in enumerate(element['header']):
            cell = table.rows[0].cells[i]
            cell.text = cell_text
            self._apply_table_cell_format(cell, table_format, is_header=True)

        # 填充表格数据
        for row_idx, row_data in enumerate(element['body'], start=1):
            for col_idx, cell_text in enumerate(row_data):
                if col_idx < num_cols:  # 防止列数不匹配
                    cell = table.rows[row_idx].cells[col_idx]
                    cell.text = cell_text
                    self._apply_table_cell_format(cell, table_format, is_header=False)

    def _add_image(self, element: Dict[str, Any]):
        """添加图片

        Args:
            element: 图片元素
        """
        try:
            src = element['src']

            # 判断是 URL 还是本地路径
            if src.startswith('http://') or src.startswith('https://'):
                # 下载网络图片
                response = requests.get(src, timeout=10)
                response.raise_for_status()
                image_stream = BytesIO(response.content)
                self.doc.add_picture(image_stream, width=Inches(5))
            else:
                # 本地文件
                if Path(src).exists():
                    self.doc.add_picture(src, width=Inches(5))
                else:
                    # 文件不存在，添加占位文本
                    self._add_placeholder(f"[图片不存在: {element.get('alt', src)}]")

        except Exception as e:
            # 图片加载失败，添加占位文本
            self._add_placeholder(f"[图片加载失败: {element.get('alt', 'N/A')}]")

    def _add_placeholder(self, text: str):
        """添加占位文本

        Args:
            text: 占位文本
        """
        para = self.doc.add_paragraph(text)
        for run in para.runs:
            run.font.color.rgb = RGBColor(255, 0, 0)  # 红色
            run.font.italic = True
            # 确保占位文本也使用正确的中文字体
            run.font.name = '宋体'
            run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')

    def _apply_paragraph_format(self, paragraph, format_dict: Dict[str, Any]):
        """应用段落格式

        Args:
            paragraph: python-docx 段落对象
            format_dict: 格式参数字典
        """
        # 应用段落级别的格式
        if format_dict.get('alignment') is not None:
            paragraph.alignment = format_dict['alignment']
        if format_dict.get('space_before') is not None:
            paragraph.paragraph_format.space_before = format_dict['space_before']
        if format_dict.get('space_after') is not None:
            paragraph.paragraph_format.space_after = format_dict['space_after']
        if format_dict.get('line_spacing') is not None:
            paragraph.paragraph_format.line_spacing = format_dict['line_spacing']
        if format_dict.get('first_line_indent') is not None:
            paragraph.paragraph_format.first_line_indent = format_dict['first_line_indent']
        if format_dict.get('left_indent') is not None:
            paragraph.paragraph_format.left_indent = format_dict['left_indent']
        if format_dict.get('right_indent') is not None:
            paragraph.paragraph_format.right_indent = format_dict['right_indent']

        # 应用运行级别的格式（字体）
        for run in paragraph.runs:
            # 分别处理西文字体和东亚字体
            font_name = format_dict.get('font_name')
            east_asia_font = format_dict.get('east_asia_font')

            # 设置西文字体
            if font_name:
                run.font.name = font_name

            # 设置东亚字体（中文字体）
            if east_asia_font:
                run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia_font)
            elif font_name:
                # 如果没有单独的东亚字体，使用西文字体作为东亚字体
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)

            if format_dict.get('font_size') is not None:
                run.font.size = format_dict['font_size']
            if format_dict.get('bold') is not None:
                run.font.bold = format_dict['bold']
            if format_dict.get('italic') is not None:
                run.font.italic = format_dict['italic']
            if format_dict.get('color') is not None:
                run.font.color.rgb = format_dict['color']

    def _apply_table_cell_margins(self, table, table_format: Dict[str, Any]):
        """应用表格单元格内边距

        Args:
            table: 表格对象
            table_format: 表格格式字典
        """
        cell_margins = table_format.get('cell_margins', {})

        # 如果没有设置内边距，使用默认值（左右5.4pt = 108 twips）
        if not cell_margins:
            cell_margins = {
                'left': 108,   # 5.4pt
                'right': 108,  # 5.4pt
            }

        # 获取表格元素
        tbl = table._element
        tblPr = tbl.find(qn('w:tblPr'))

        if tblPr is None:
            from docx.oxml import OxmlElement
            tblPr = OxmlElement('w:tblPr')
            tbl.insert(0, tblPr)

        # 设置单元格内边距
        tblCellMar = tblPr.find(qn('w:tblCellMar'))
        if tblCellMar is None:
            from docx.oxml import OxmlElement
            tblCellMar = OxmlElement('w:tblCellMar')
            tblPr.append(tblCellMar)

        # 设置各个方向的内边距
        for side, value in cell_margins.items():
            elem = tblCellMar.find(qn(f'w:{side}'))
            if elem is None:
                from docx.oxml import OxmlElement
                elem = OxmlElement(f'w:{side}')
                tblCellMar.append(elem)
            elem.set(qn('w:w'), str(value))
            elem.set(qn('w:type'), 'dxa')

    def _adjust_table_column_widths(self, table, element: Dict[str, Any]):
        """根据模版列宽比例调整表格列宽

        Args:
            table: 表格对象
            element: 表格元素（包含 header 和 body）
        """
        from docx.shared import Cm

        num_cols = len(element['header']) if element['header'] else 0
        if num_cols == 0:
            return

        # 获取模版中的列宽比例
        table_format = self.format_styles.get('table', {})
        column_width_ratios = table_format.get('column_width_ratios', [])

        # 如果模版有列宽比例，使用模版比例
        if column_width_ratios and len(column_width_ratios) >= num_cols:
            # 页面可用宽度（根据模版表格总宽度，约17.38cm）
            # 使用模版的实际总宽度
            column_widths = table_format.get('column_widths', [])
            if column_widths and len(column_widths) >= num_cols:
                # 直接使用模版中的列宽
                for i, col in enumerate(table.columns):
                    if i < len(column_widths):
                        col.width = column_widths[i]
        # 否则使用内容长度计算（原有逻辑）
        else:
            # 计算每列的最大文本长度
            col_max_lengths = [0] * num_cols

            # 统计表头长度
            for i, header_text in enumerate(element['header']):
                col_max_lengths[i] = max(col_max_lengths[i], len(header_text))

            # 统计数据行长度
            for row_data in element['body']:
                for i, cell_text in enumerate(row_data):
                    if i < num_cols:
                        col_max_lengths[i] = max(col_max_lengths[i], len(cell_text))

            # 计算总长度
            total_length = sum(col_max_lengths)
            if total_length == 0:
                return

            # 页面可用宽度（A4纸宽21cm，减去左右页边距各2.5cm = 16cm）
            available_width = Cm(16)

            # 根据内容长度按比例分配列宽
            for i, col in enumerate(table.columns):
                # 计算该列应占的比例
                ratio = col_max_lengths[i] / total_length
                # 设置最小宽度为2cm，避免列太窄
                col.width = max(Cm(2), int(available_width * ratio))

    def _apply_table_cell_format(self, cell, table_format: Dict[str, Any], is_header: bool = False):
        """应用表格单元格格式

        Args:
            cell: 表格单元格对象
            table_format: 表格格式字典
            is_header: 是否为表头
        """
        if not cell.paragraphs:
            return

        para = cell.paragraphs[0]
        cell_text = para.text  # 保存文本内容

        # 从模版中获取对齐方式
        if is_header:
            # 应用表头对齐方式（默认居中）
            header_alignment = table_format.get('header_alignment')
            if header_alignment is not None:
                para.alignment = header_alignment
            # 应用表头首行缩进（只有非0时才设置）
            header_indent = table_format.get('header_indent')
            if header_indent is not None and header_indent != 0:
                para.paragraph_format.first_line_indent = header_indent
        else:
            # 应用数据行对齐方式（默认左对齐）
            data_alignment = table_format.get('data_alignment')
            if data_alignment is not None:
                para.alignment = data_alignment
            # 应用数据行首行缩进（只有非0时才设置）
            data_indent = table_format.get('data_indent')
            if data_indent is not None and data_indent != 0:
                para.paragraph_format.first_line_indent = data_indent

        # 清除现有的 runs 并重新创建带格式的 run
        para.clear()
        run = para.add_run(cell_text)

        # 应用字体格式
        font_name = table_format.get('cell_font_name')
        east_asia_font = table_format.get('cell_east_asia_font')

        # 设置西文字体
        if font_name:
            run.font.name = font_name

        # 设置东亚字体（中文字体）
        if east_asia_font:
            run._element.rPr.rFonts.set(qn('w:eastAsia'), east_asia_font)
        elif font_name:
            # 如果没有单独的东亚字体，使用西文字体作为东亚字体
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        if table_format.get('cell_font_size') is not None:
            run.font.size = table_format['cell_font_size']

        # 从模版中获取加粗设置
        if is_header:
            header_bold = table_format.get('header_bold')
            if header_bold is not None:
                run.font.bold = header_bold
        else:
            data_bold = table_format.get('data_bold')
            if data_bold is not None:
                run.font.bold = data_bold

    def save(self):
        """保存文档"""
        self.doc.save(self.output_path)
